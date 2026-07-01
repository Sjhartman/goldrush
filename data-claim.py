"""
data-claim.py -- IRB Data Request Validator

Compares a collaborator's data request against their IRB protocol.
Outputs a structured audit JSON (input to excavator/orchestrator.py) and
a human-readable Markdown report.

Supported formats (IRB and request docs): PDF, DOCX, HTML, HTM, TXT, MD, CSV, TSV

Usage:
  venv/bin/python data-claim.py <irb_document> <request_document> [options]

Options:
  --clarification <path_or_text>
                    Follow-up clarifications resolving ambiguous wording in the
                    data request. Can be a file path or inline text. DOCX
                    clarifications with ICD tables are merged with the request's
                    tables -- clarification entries supersede request entries for
                    the same diagnosis label.
  --icd-file <path> ICD-10 reference CSV. Defaults to the newest
                    ICD10_codes/icd10cm-oncology-*.csv found next to this script.
  --override-i2db   Waive the I2DB/RDC authorization check at reviewer discretion.
"""

import anthropic
import base64
import json
import os
import re
import sys
import time
from datetime import datetime, timezone
from pathlib import Path


SYSTEM_PROMPT = """You are a research compliance assistant helping validate data requests against IRB protocols.
The underlying dataset uses ICD-10 codes exclusively.

You will be given:
1. An IRB protocol document
2. A collaborator's data request
3. (Optional) Follow-up clarifications from the requestor that resolve ambiguous wording in
   the data request. Treat these as authoritative for interpreting the request's intent, but
   all requested scope must still fall within what the IRB protocol explicitly permits.

ICD CODE HANDLING:
- The underlying data uses ICD-10 codes only.
- Requestors may specify diagnoses using disease names, ICD-9 codes, ICD-10 codes, or a mix.
- For any diagnosis mentioned, identify whether it was specified as a disease name, ICD-9, ICD-10, or mixed.
- If BOTH ICD-9 and ICD-10 codes are present for the same diagnosis, use the ICD-10 codes and
  ignore the ICD-9 codes. Note in mapping_notes that ICD-9 codes were present but superseded.
- If ONLY ICD-9 codes are present with no ICD-10 equivalent provided, do NOT attempt to translate
  them. Set requires_clarification to true and note that ICD-10 codes must be supplied by the
  requestor before this item can be processed.
- If only a disease name is given with no codes at all, set requires_clarification to true
  unless the formal aims individually enumerate all the cancer types implied by that label --
  in that case, note the constituent mapping and set requires_clarification to false.
- If the cohort is primarily identified by molecular or genomic test results (e.g., NGS panel
  findings, gene alterations, Tempus/Foundation/CARIS reports) rather than by diagnosis code,
  ICD-10 codes are still required as basic inclusion criteria for data filtering. If the cohort
  description references an organ site or disease category (e.g., "GI or HPB pathology specimens",
  "lung cancer", "breast tumor"), set code_system_used to "DISEASE_NAME" and requires_clarification
  to true so codes are resolved from the ICD-10 reference table. Use the organ site or disease
  category as the as_specified label. Do NOT set code_system_used to "NONE" when an organ site
  or disease type is named in the request -- NONE is only correct when absolutely no anatomical
  or diagnostic information is provided.
  Do NOT add an icd_flags entry for "no codes provided" -- their absence in the request is
  expected and not an issue. Note in mapping_notes that ICD filtering is supplementary to
  the molecular cohort definition and the codes require verification before use.
  Auto-resolved mapping accuracy should still be flagged as usual.
- Always report the final ICD-10 codes that would be used for data filtering.
- Parent ICD-10 codes (e.g. C00, Z15, K50) cover all subcodes by ICD-10 convention.
  Do NOT flag a parent code as incomplete or abbreviated -- C00 validly covers C00.0,
  C00.1, etc. Only flag a code list as incomplete if specific subcodes are genuinely
  needed that the parent does not cover.
- When referencing ICD-10 code lists in your output -- especially exclusion criteria --
  reproduce every individual code exactly as provided. Never compress a list of
  individual codes into a range notation (e.g. do NOT write "C00-C14"; write each
  code separately). Abbreviation loses information and will cause incorrect SQL filters.

ICD FILTERING ARCHITECTURE:
- ICD-10 codes provide initial diagnosis-based cohort filtering only.
- Clinical attributes encodable via Z codes (e.g. hormone receptor status via Z17.x) are resolved
  as qualifier codes alongside the primary diagnosis codes -- do not ask the requestor to supply these.
- Metastatic/advanced disease is expressed using secondary malignant neoplasm codes (C77.x, C78.x,
  C79.x) as qualifier codes when the IRB specifies a metastatic population.
- The final population filter should be expressed as explicit combination logic
  (e.g. "C50.x AND (Z17.0 OR Z17.21) AND (C77.x OR C78.x OR C79.x)").

EXCLUSION CRITERIA -- populate both sources and cross-check:
- irb_summary.exclusion_criteria: draw ONLY from the IRB protocol's formal eligibility
  or procedures sections. Do NOT include standard IRB template vulnerability categories
  (e.g. minors, prisoners, pregnant women from boilerplate section 2.x sections) unless they
  are explicitly cited as study-specific exclusions in the formal aims or eligibility text.
- request_summary.exclusion_criteria: draw ONLY from the data request document. These are
  patient groups or record types the requestor explicitly wants excluded.

SCOPE POLICY -- the IRB defines the maximum allowable scope, not a minimum:
- The requestor is always permitted to request LESS than what the IRB approves.
  If the data request excludes a population the IRB would allow, that is the requestor's
  choice and must be respected. Apply the exclusion and note in population_filter_check
  that the requestor is voluntarily limiting their scope below what the IRB permits.
- Do NOT flag a request exclusion as ambiguous simply because the IRB does not mention it.
  An exclusion is only problematic if it conflicts with something the IRB explicitly
  REQUIRES to be included (rare), or if the scope of the exclusion is genuinely unclear.
- Only flag cross_document_conflicts when a request exclusion would REMOVE patients the
  IRB mandates must be in scope, not merely because the exclusion is more restrictive.

CONSISTENCY CHECKING -- populate the consistency_checks fields:
1. Pre-computed ICD conflicts will be provided in the user message (if any were found). Each one
   MUST be included verbatim in consistency_checks.internal_request_conflicts, with a plain-English
   explanation of what it means for the data team and how to resolve it.
2. Also flag any other intra-document contradictions you find: conflicting age ranges, population
   scopes, or criteria named in both include and exclude contexts.
3. Cross-document conflicts: IRB exclusion criteria violated by the request, or formal IRB
   inclusion criteria missing from the request's population filters.

Your job is to produce a structured JSON report with exactly these fields:

{
  "irb_summary": {
    "protocol_number": "string or null",
    "pi_name": "string or null",
    "approved_personnel": ["list of names/roles"],
    "study_population": "description of who is included",
    "inclusion_criteria": ["list of criteria from the IRB formal eligibility section"],
    "exclusion_criteria": ["study-specific exclusion criteria from the IRB formal eligibility section only -- NOT boilerplate vulnerability categories (minors, prisoners, pregnant women) from standard IRB template sections"],
    "approved_data_elements": ["list of data types approved"],
    "approval_expiry": "date string or null"
  },
  "request_summary": {
    "requester": "string or null",
    "data_elements_requested": ["list of what they want"],
    "population_filters": ["inclusion filters from the data request: age ranges, diagnoses, date ranges, etc."],
    "exclusion_criteria": ["patient groups or records explicitly excluded in the data request -- draw ONLY from the data request document, not from the IRB"],
    "intended_use": "string or null"
  },
  "icd_code_analysis": {
    "diagnoses_requested": [
      {
        "as_specified": "exactly how the requestor wrote it",
        "code_system_used": "ICD-9 | ICD-10 | DISEASE_NAME | MIXED | NONE",
        "icd9_codes": ["codes if ICD-9 was used, else empty"],
        "icd10_codes": ["final ICD-10 codes to use for filtering"],
        "qualifier_condition": "non-ICD filter condition attached to this code (e.g. 'age < 40', 'prior to index diagnosis') -- null if unconditional",
        "mapping_notes": "any ambiguity, one-to-many mappings, or assumptions made",
        "requires_clarification": true
      }
    ],
    "exclusion_diagnoses": [
      {
        "as_specified": "exactly how the requestor wrote this exclusion",
        "icd10_codes": ["ICD-10 codes to exclude"],
        "icd9_codes": ["ICD-9 codes if present, else empty"],
        "qualifier_condition": "non-ICD filter condition attached to this exclusion (e.g. 'age < 40', 'prior to index diagnosis date') -- null if unconditional",
        "notes": "any ambiguity or verification needed"
      }
    ],
    "icd_flags": ["list of any ICD translation issues requiring human review"]
  },
  "validation": {
    "overall_status": "APPROVED | DENIED | NEEDS_REVIEW",
    "approved_elements": [
      {
        "element": "data element name",
        "rationale": "specific IRB language supporting this"
      }
    ],
    "denied_elements": [
      {
        "element": "data element name",
        "rationale": "why this is outside IRB scope"
      }
    ],
    "ambiguous_elements": [
      {
        "element": "data element name",
        "reason": "why this is unclear",
        "irb_language": "relevant IRB text",
        "recommendation": "what reviewer should check"
      }
    ],
    "personnel_check": {
      "status": "APPROVED | DENIED | NOT_SPECIFIED",
      "notes": "string"
    },
    "population_filter_check": [
      {
        "filter": "filter description",
        "status": "WITHIN_SCOPE | OUTSIDE_SCOPE | AMBIGUOUS",
        "notes": "string"
      }
    ]
  },
  "consistency_checks": {
    "internal_request_conflicts": [
      {
        "description": "what conflicts -- be specific about which codes or criteria appear in both lists",
        "conflicting_items": ["item or code range A", "item or code range B"],
        "recommendation": "how the data team should resolve this -- 1-2 sentences"
      }
    ],
    "cross_document_conflicts": [
      {
        "description": "how the IRB and request contradict each other",
        "irb_basis": "relevant IRB text",
        "request_basis": "relevant request text",
        "recommendation": "how to resolve -- 1-2 sentences"
      }
    ]
  },
  "available_but_not_requested": [
    "Data elements or patient populations the IRB explicitly approves but the requestor did not ask for. Each entry should name the element and cite the IRB section. This helps reviewers understand what scope was left unused and could be added to a future request without an amendment."
  ],
  "reviewer_notes": "Overall summary for human reviewer, highlighting key concerns including any ICD translation issues"
}

DATA SOURCE DATE RANGE RULES -- apply these strictly:
Some IRB protocols list multiple data sources (institutions, registries, or databases) each
with its own applicable date range. Follow these rules when reading such sections:

1. OLD vs. NEW subsections: If a data source section (e.g. section 1.24) contains an "old" and a
   "new" subsection (or similarly versioned text), ignore the "old" section entirely. Extract
   date ranges and source descriptions ONLY from the "new" (current) section.

2. Multi-institution protocols: When multiple institutions are listed with different date
   ranges, identify which institution(s) are the actual data providers for this request.
   For WashU / Siteman Cancer Center studies, the primary EHR sources are Barnes Jewish
   Hospital (BJC/BJH) and Washington University School of Medicine (WU/WUSM). Prioritize
   those institutions' date ranges as the operative constraints.

3. Date range as a population filter: The applicable date range for the relevant institution
   IS a population-scope constraint. Include it as a population_filter_check entry with
   status WITHIN_SCOPE (if the request respects it) or OUTSIDE_SCOPE (if the request asks
   for data outside that range). If the request does not specify a date range but the IRB
   restricts one, note the IRB constraint and flag it for the data team to enforce.

4. Cite the specific section and institution: In notes and rationale, name the section
   (e.g. "section 1.24 New Value") and the institution (e.g. "Barnes Jewish Hospital records") and
   the exact date range (e.g. "01/01/1998 - 12/31/2022").

GENOMIC DATA SOURCE AUTHORIZATION -- apply when Tempus or NGS data is requested:
- When the data request asks for Tempus NGS/genomic profiling data (or any molecular profiling
  results from the WashU/BJC Healthcare system), the IRB protocol MUST explicitly authorize
  acquisition from the Institute for Informatics, Data Science, and Biostatistics (I2DB)
  Research Data Core (RDC) Repository (IRB ID: 201607071, PI: Albert Lai).
- Search the IRB document for any reference to: "I2DB", "Institute for Informatics",
  "Research Data Core", "RDC", "IRB 201607071", or "Albert Lai".
- If none of these are found: mark the Tempus/genomic data element as DENIED in
  denied_elements, and add a cross_document_conflicts entry stating that WashU/BJC Tempus
  genomic profiling data is sourced from the I2DB RDC Repository (IRB 201607071, PI: Albert Lai)
  and this authorization is absent from the IRB protocol. Set overall_status to DENIED.
- If the IRB does reference the RDC/I2DB authorization: the genomic data element may proceed
  to be evaluated on its own merits (approved, ambiguous, or denied) based on the IRB scope.

PERSONNEL CHECK RULES -- apply these strictly:
- Set status to NOT_SPECIFIED when the data request does not name any specific personnel
  (no individual names, no requester listed, or only a PI name carried over from the IRB).
  Do NOT mark personnel APPROVED simply because the IRB has an approved roster -- approval
  only applies when the request explicitly names people AND all of them are on the roster.
- Set status to APPROVED only when the request names specific individuals AND every one of
  them appears on the IRB's approved personnel list.
- Set status to DENIED when the request names at least one person not on the IRB roster.
- The notes field must state who was named in the request (or "No personnel named in data
  request") and what was found on the IRB roster.

Rules:
- Be conservative. When in doubt, mark as AMBIGUOUS rather than APPROVED.
- Always cite specific IRB language when approving or denying elements.
- Never infer approval -- it must be explicitly supported by IRB text.
- If the IRB document is unclear or incomplete, note this prominently.
- When a data request form contains an "Others" or open-ended checkbox, look for clarifying
  text in the immediately following section, adjacent free-text fields, or nearby descriptive
  text before marking it as ambiguous. If the form explains what "Others" means in a subsequent
  section, use that text to populate the data element -- do NOT flag it as ambiguous.
- For ICD codes: never silently assume a mapping is correct if there is any ambiguity.
- When identifying study population, inclusion criteria, and exclusion criteria, draw ONLY from
  the protocol's formal eligibility, aims, and procedures sections. Do NOT infer or apply
  qualifiers from background sections, introduction, the study title, or referenced literature.
  A qualifier must be explicitly stated as a requirement in the formal protocol text to be
  treated as a population filter criterion. Qualifiers that appear only in background or title
  context are informational and must not be applied as hard filters.
- Always populate consistency_checks by comparing inclusion vs. exclusion ICD code lists in the
  request. If no conflicts exist, return empty arrays. Do not omit the field.
- For icd_code_analysis.exclusion_diagnoses: populate one entry per distinct exclusion row in the
  data request's exclusion ICD table. If a qualifier condition is attached to the ENTIRE row
  (e.g. "prior to index diagnosis"), capture it verbatim in qualifier_condition. If a qualifier
  condition is attached to a SPECIFIC CODE within a row -- indicated by parenthetical text
  immediately after the code, such as "D12.6 (combine with age <40)" in a list of other codes --
  create a SEPARATE exclusion_diagnoses entry for that code alone with the parenthetical text
  captured verbatim in qualifier_condition, and omit that code from the parent row's entry.
  Example: a cell containing "Q85.81, D12.6 (combine with age <40), Z80.0" must produce TWO
  exclusion_diagnoses entries: one for [Q85.81, Z80.0] with qualifier_condition null, and one
  for [D12.6] with qualifier_condition "combine with age <40".
  Unconditional codes get qualifier_condition: null. Never strip or paraphrase qualifiers.
- Be concise throughout. Each rationale, reason, recommendation, and notes field must be 1-2
  sentences maximum. Do not repeat information already captured in another field. The
  reviewer_notes field must be a tight bullet list -- one line per issue, no prose paragraphs.
  Omit any field that would be empty or redundant.
- Consolidate entries that share the same rationale, reason, or recommendation. If multiple
  approved elements are approved for the same IRB reason, write one entry listing all their names
  (e.g. "Colorectal, pancreatic, liver, and lung cancer cases"). If multiple denied or ambiguous
  elements share the same reason and recommendation, write one combined entry. Apply this to
  approved_elements, denied_elements, ambiguous_elements, and population_filter_check rows.
- Populate available_but_not_requested with IRB-approved elements or populations that
  the request did not ask for. Compare irb_summary.approved_data_elements against
  request_summary.data_elements_requested. Each entry must cite the IRB section and name
  the element concisely. If everything approved was requested, write an empty list.
- Call the submit_validation tool with your complete result. Do not output any text outside the tool call."""


SUPPORTED_FORMATS = ".pdf, .docx, .html, .htm, .txt, .md, .csv, .tsv"

VALIDATION_TOOL = {
    "name": "submit_validation",
    "description": "Submit the complete IRB validation result.",
    "input_schema": {
        "type": "object",
        "properties": {
            "irb_summary":                 {"type": "object"},
            "request_summary":             {"type": "object"},
            "icd_code_analysis":           {"type": "object"},
            "validation":                  {"type": "object"},
            "consistency_checks":          {"type": "object"},
            "available_but_not_requested": {"type": "array", "items": {"type": "string"}},
            "reviewer_notes":              {"type": "string"},
        },
        "required": [
            "irb_summary", "request_summary", "icd_code_analysis",
            "validation", "consistency_checks", "available_but_not_requested",
            "reviewer_notes",
        ],
    },
}

ICD_LOOKUP_TOOL = {
    "name": "submit_icd_lookup",
    "description": "Submit the ICD-10 code lookup result.",
    "input_schema": {
        "type": "object",
        "properties": {
            "lookups": {"type": "array"},
        },
        "required": ["lookups"],
    },
}


def load_document(file_path: str, label: str) -> tuple:
    """
    Load a document into API content blocks.
    Returns (content_blocks, format_description).
    """
    path   = Path(file_path)
    suffix = path.suffix.lower()

    if suffix == ".pdf":
        with open(file_path, "rb") as f:
            b64 = base64.standard_b64encode(f.read()).decode("utf-8")
        content = [{
            "type":   "document",
            "source": {"type": "base64", "media_type": "application/pdf", "data": b64},
            "title":  label,
        }]
        return content, "PDF"

    elif suffix == ".docx":
        try:
            from docx import Document
        except ImportError:
            raise ImportError("pip install python-docx  # required for .docx support")
        doc  = Document(file_path)
        text = "\n\n".join(p.text for p in doc.paragraphs if p.text.strip())
        return [{"type": "text", "text": f"{label} (from DOCX):\n\n{text}"}], "DOCX"

    elif suffix in (".html", ".htm"):
        try:
            from bs4 import BeautifulSoup
        except ImportError:
            raise ImportError("pip install beautifulsoup4  # required for HTML support")
        with open(file_path, "r", encoding="utf-8", errors="replace") as f:
            soup = BeautifulSoup(f.read(), "html.parser")
        for tag in soup(["script", "style", "nav", "footer"]):
            tag.decompose()
        text = soup.get_text(separator="\n", strip=True)
        return [{"type": "text", "text": f"{label} (from HTML):\n\n{text}"}], "HTML"

    elif suffix in (".txt", ".md", ".csv", ".tsv"):
        with open(file_path, "r", encoding="utf-8", errors="replace") as f:
            text = f.read()
        fmt = suffix.upper().strip(".")
        return [{"type": "text", "text": f"{label} (from {fmt}):\n\n{text}"}], fmt

    else:
        raise ValueError(
            f"Unsupported format: '{suffix}' for {file_path}. "
            f"Supported formats: {SUPPORTED_FORMATS}"
        )


ICD_LOOKUP_SYSTEM = """You are an ICD-10 code lookup assistant for clinical research compliance.
Given one or more disease names and IRB population criteria, generate a complete combination filter
using the provided ICD-10 reference.

Return ONLY valid JSON with exactly this structure:
{
  "lookups": [
    {
      "as_specified": "exactly as provided in the input list",
      "primary_codes": [
        {"code": "C50.011", "description": "description from reference"}
      ],
      "qualifier_codes": [
        {
          "code": "Z17.0",
          "description": "Estrogen receptor positive status [ER+]",
          "role": "hormone_receptor_positive"
        }
      ],
      "combination_logic": "C50.x AND (Z17.0 OR Z17.21) AND (C77.x OR C78.x OR C79.x)",
      "metastatic_caveat": "string if metastatic qualifier was applied, else null",
      "lookup_notes": "any ambiguity, broad mappings, or assumptions",
      "requires_human_review": true
    }
  ]
}

Rules:
- primary_codes: all ICD-10 C/D diagnosis codes for the disease name. Include ALL subcodes.
- qualifier_codes: Z codes for clinical attributes required by the IRB population criteria
  (e.g. Z17.0 for ER+, Z17.21 for PR+, Z17.31 for HER2+). Only include qualifiers explicitly
  required by the IRB criteria passed with the request.
- For metastatic or advanced disease: include C77.x, C78.x, and C79.x ranges as qualifier codes
  with role "metastatic". Set metastatic_caveat to null.
- combination_logic: express the full filter as explicit AND/OR logic using code ranges.
- If a disease maps to a broad category, include all subcodes and set requires_human_review to true.
- Keep lookup_notes and metastatic_caveat to 1-2 sentences each. Do not repeat information
  already expressed in combination_logic.
- Call the submit_icd_lookup tool with your complete result."""


def resolve_icd_codes(disease_names: list, icd_file_path: str, irb_context: dict = None) -> tuple:
    """
    Maps disease names to ICD-10 combination filters using a reference file.
    Returns (lookup_result, usage).
    """
    client = anthropic.Anthropic()
    model  = os.environ.get("ANTHROPIC_MODEL", "claude-sonnet-4-6")

    icd_content, _ = load_document(icd_file_path, "ICD-10 Code Reference")
    icd_content[-1]["cache_control"] = {"type": "ephemeral"}

    irb_context_text = ""
    if irb_context:
        study_pop = irb_context.get("study_population", "")
        inclusion = irb_context.get("inclusion_criteria", [])
        if study_pop or inclusion:
            irb_context_text = (
                f"\n\nIRB-approved study population:\n{study_pop}\n"
                f"Inclusion criteria:\n"
                + "\n".join(f"  - {c}" for c in inclusion)
            )

    names_list  = "\n".join(f"- {name}" for name in disease_names)
    query_block = {
        "type": "text",
        "text": (
            f"Generate ICD-10 combination filters for each of the following disease names "
            f"using the ICD-10 Code Reference above.{irb_context_text}\n\n"
            f"Disease names:\n{names_list}"
        ),
    }

    message = _api_call_with_retry(lambda: client.messages.create(
        model=model,
        max_tokens=8192,
        system=ICD_LOOKUP_SYSTEM,
        tools=[ICD_LOOKUP_TOOL],
        tool_choice={"type": "tool", "name": "submit_icd_lookup"},
        messages=[{"role": "user", "content": icd_content + [query_block]}],
    ))

    tool_blocks = [b for b in message.content if b.type == "tool_use"]
    if not tool_blocks:
        raise RuntimeError("ICD lookup returned no tool_use block.")
    return tool_blocks[0].input, message.usage


def apply_icd_resolution(result: dict, lookup_result: dict) -> dict:
    lookup_map = {e["as_specified"]: e for e in lookup_result.get("lookups", [])}
    diagnoses  = result.get("icd_code_analysis", {}).get("diagnoses_requested", [])

    needs_review_names = []
    metastatic_caveats = []
    missing_names      = []

    for diagnosis in diagnoses:
        if not (diagnosis.get("requires_clarification")
                and diagnosis.get("code_system_used") == "DISEASE_NAME"):
            continue
        entry = lookup_map.get(diagnosis["as_specified"])
        if not entry:
            continue

        primary_codes      = [c["code"] for c in entry.get("primary_codes", [])]
        qualifier_codes    = [c["code"] for c in entry.get("qualifier_codes", [])]
        combination_logic  = entry.get("combination_logic", "")
        metastatic_caveat  = entry.get("metastatic_caveat")
        notes              = entry.get("lookup_notes", "")

        if primary_codes:
            diagnosis["icd10_codes"] = primary_codes
            if qualifier_codes:
                diagnosis["qualifier_codes"] = qualifier_codes
            if combination_logic:
                diagnosis["combination_logic"] = combination_logic
            diagnosis["mapping_notes"] = (
                f"[Auto-resolved from ICD-10 reference -- verify before use] {notes}".strip()
            )
            if not entry.get("requires_human_review", True):
                diagnosis["requires_clarification"] = False
            else:
                needs_review_names.append(diagnosis["as_specified"])
            if metastatic_caveat and metastatic_caveat not in metastatic_caveats:
                metastatic_caveats.append(metastatic_caveat)
        else:
            missing_names.append(diagnosis["as_specified"])

    extra_flags = []
    if needs_review_names:
        names = ", ".join(f"'{n}'" for n in needs_review_names)
        extra_flags.append(
            f"Auto-resolved ICD-10 mappings for {names} require human verification "
            f"before use -- codes are standard but may need refinement for this study."
        )
    for caveat in metastatic_caveats:
        extra_flags.append(f"Metastatic qualifier note: {caveat}")
    if missing_names:
        names = ", ".join(f"'{n}'" for n in missing_names)
        extra_flags.append(
            f"No ICD-10 codes found in reference file for {names} -- "
            f"must be supplied by the requestor."
        )

    if extra_flags:
        result["icd_code_analysis"].setdefault("icd_flags", []).extend(extra_flags)
    return result


def _extract_icd_codes_from_cell(text: str) -> list:
    codes = []
    for part in re.split(r"[,\n]", text):
        part = part.strip()
        m    = re.match(r"^([A-Z]\d{1,2}(?:\.\d{0,4}[A-Z0-9]*)?)(?:[\s\-\(]|$)", part)
        if m:
            codes.append(m.group(1))
    return codes


def _extract_docx_icd_table(file_path: str) -> dict | None:
    try:
        from docx import Document
        doc = Document(file_path)
    except Exception:
        return None

    result: dict         = {"inclusions": [], "exclusions": []}
    current_section      = None

    for table in doc.tables:
        for row in table.rows:
            cells     = [cell.text.strip() for cell in row.cells]
            non_empty = [c for c in cells if c]
            if not non_empty:
                continue

            first_upper = non_empty[0].upper()
            if len(set(non_empty)) == 1:
                if "EXCLUSION" in first_upper:
                    current_section = "exclusions"
                elif "INCLUSION" in first_upper:
                    current_section = "inclusions"
                continue

            if current_section is None or len(cells) < 2:
                continue

            label = cells[0]
            if not label or re.search(r"^ICD|^SOURCE", label, re.I):
                continue

            icd10_codes = _extract_icd_codes_from_cell(cells[1] if len(cells) > 1 else "")
            icd9_raw    = cells[2] if len(cells) > 2 else ""
            icd9_codes  = [
                p.strip() for p in re.split(r"[,\n]", icd9_raw)
                if re.match(r"^\d{3}(\.\d+)?$", p.strip())
            ] if icd9_raw and not re.search(r"N/A|not captured|predates", icd9_raw, re.I) else []

            if icd10_codes or icd9_codes:
                result[current_section].append((label, icd10_codes, icd9_codes))

    return result if (result["inclusions"] or result["exclusions"]) else None


def _find_icd_conflicts(icd_data: dict) -> list:
    conflicts = []
    for inc_label, inc_codes, _inc9 in icd_data["inclusions"]:
        inc_roots = {c.split(".")[0] for c in inc_codes}
        for exc_label, exc_codes, _exc9 in icd_data["exclusions"]:
            exc_roots = {c.split(".")[0] for c in exc_codes}
            overlap   = sorted(inc_roots & exc_roots)
            if overlap:
                conflicts.append({
                    "inclusion_label": inc_label,
                    "exclusion_label": exc_label,
                    "overlapping_codes": overlap,
                })
    return conflicts


def _build_conflict_entries(conflicts: list) -> list:
    from collections import defaultdict
    by_exclusion: dict = defaultdict(list)
    for c in conflicts:
        by_exclusion[c["exclusion_label"]].append(c)

    entries = []
    for exc_label, group in by_exclusion.items():
        all_codes  = sorted({code for c in group for code in c["overlapping_codes"]})
        items_list = [
            f"Inclusion -- {c['inclusion_label']}: {', '.join(c['overlapping_codes'])}"
            for c in group
        ] + [f"Exclusion -- {exc_label}"]

        if len(group) == 1:
            c          = group[0]
            codes_str  = ", ".join(c["overlapping_codes"])
            description = (
                f"Inclusion '{c['inclusion_label']}' contains ICD root codes ({codes_str}) "
                f"that also appear in the exclusion criterion '{exc_label}'. "
                f"A patient with this diagnosis would be simultaneously included and excluded."
            )
        else:
            inc_summary = "; ".join(
                f"{c['inclusion_label']} ({', '.join(c['overlapping_codes'])})" for c in group
            )
            description = (
                f"Multiple inclusion diagnoses share ICD root codes with the exclusion "
                f"criterion '{exc_label}': {inc_summary}. "
                f"As written, patients with any of these study diagnoses would be simultaneously "
                f"included and excluded -- the exclusion criterion must be refined before any "
                f"cohort filter can be applied."
            )

        entries.append({
            "description":      description,
            "conflicting_items": items_list,
            "recommendation": (
                f"Refine '{exc_label}' to apply only to prior or concurrent malignancies "
                f"OTHER THAN the index cancer under study. Confirm with the requestor which "
                f"patients are truly meant to be excluded (e.g., those with a prior different "
                f"malignancy) and update the cohort definition before applying any exclusion "
                f"filter. Affected root codes: {', '.join(all_codes)}."
            ),
        })
    return entries


def _merge_icd_tables(primary: dict, clarification: dict) -> dict:
    result = {}
    for section in ("inclusions", "exclusions"):
        by_label = {
            label: (label, icd10, icd9)
            for label, icd10, icd9 in primary.get(section, [])
        }
        for label, icd10, icd9 in clarification.get(section, []):
            by_label[label] = (label, icd10, icd9)
        result[section] = list(by_label.values())
    return result


def _api_call_with_retry(call_fn, max_retries: int = 3, base_delay: float = 5.0,
                         response_validator=None):
    for attempt in range(max_retries):
        try:
            message = call_fn()
            if not message.content:
                reason = "empty response"
            elif response_validator:
                ok, reason = response_validator(message)
                if ok:
                    return message
            else:
                return message
            if attempt == max_retries - 1:
                raise RuntimeError(
                    f"API response invalid after {max_retries} attempts: {reason}."
                )
            delay = base_delay * (2 ** attempt)
            print(f"  API response invalid ({reason}), retrying in {delay:.0f}s "
                  f"(attempt {attempt + 1}/{max_retries})...")
            time.sleep(delay)
        except (anthropic.InternalServerError, anthropic.APIConnectionError) as e:
            if attempt == max_retries - 1:
                raise
            delay = base_delay * (2 ** attempt)
            print(f"  API error ({e}), retrying in {delay:.0f}s (attempt {attempt + 1}/{max_retries})...")
            time.sleep(delay)


def validate_request(irb_path: str, request_arg: str, clarification_arg: str = None,
                     override_i2db: bool = False) -> tuple:
    client = anthropic.Anthropic()
    model  = os.environ.get("ANTHROPIC_MODEL", "claude-sonnet-4-6")

    irb_content, irb_format = load_document(irb_path, "IRB Protocol Document")
    print(f"  IRB document:  {irb_format}")

    request_path = Path(request_arg)
    if request_path.exists():
        request_content, request_format = load_document(request_arg, "Data Request")
        print(f"  Data request:  {request_format}")
    else:
        request_content = [{"type": "text", "text": f"Data Request:\n\n{request_arg}"}]
        request_format  = "INLINE_TEXT"
        print(f"  Data request:  inline text")

    clarification_content = []
    clarification_format  = None
    clar_path             = None
    if clarification_arg:
        _clar = Path(clarification_arg)
        if _clar.exists():
            clarification_content, clarification_format = load_document(
                clarification_arg, "Request Clarifications"
            )
            clar_path = _clar
            print(f"  Clarifications: {clarification_format}")
        else:
            clarification_content = [
                {"type": "text", "text": f"Request Clarifications:\n\n{clarification_arg}"}
            ]
            clarification_format = "INLINE_TEXT"
            print(f"  Clarifications: inline text")

    icd_table = None
    if request_path.exists() and request_path.suffix.lower() == ".docx":
        icd_table = _extract_docx_icd_table(str(request_path))

    if clar_path and clar_path.suffix.lower() == ".docx":
        clar_icd_table = _extract_docx_icd_table(str(clar_path))
        if clar_icd_table:
            icd_table = _merge_icd_tables(
                icd_table or {"inclusions": [], "exclusions": []}, clar_icd_table
            )

    precomputed_entries = []
    conflict_note       = ""
    provided_codes_note = ""
    table_labels: set   = set()

    if icd_table:
        table_labels = {
            label
            for section in ("inclusions", "exclusions")
            for label, _icd10, _icd9 in icd_table.get(section, [])
        }
        conflicts = _find_icd_conflicts(icd_table)
        if conflicts:
            precomputed_entries = _build_conflict_entries(conflicts)
            lines = ["PRE-COMPUTED ICD CODE CONFLICTS (programmatically detected):"]
            for c in conflicts:
                lines.append(
                    f"  - Inclusion '{c['inclusion_label']}' vs. Exclusion "
                    f"'{c['exclusion_label']}': shared root codes {', '.join(c['overlapping_codes'])}"
                )
            lines.append(
                "\nDo NOT write entries for these in consistency_checks.internal_request_conflicts "
                "-- they will be injected programmatically. Only add entries for other "
                "intra-document conflicts NOT related to ICD inclusion/exclusion overlap."
            )
            conflict_note = "\n".join(lines)
            print(f"  Pre-computed ICD conflicts: {len(conflicts)} raw -> {len(precomputed_entries)} consolidated")

        coded_inclusions = [
            (label, icd10, icd9)
            for label, icd10, icd9 in icd_table.get("inclusions", [])
            if icd10 or icd9
        ]
        if coded_inclusions:
            code_lines = [
                "ICD CODES DETECTED IN DATA REQUEST INCLUSION TABLE (programmatically extracted):",
                "For each entry below:",
                "  - Set `as_specified` to the EXACT string shown after 'label:' -- do not modify or append anything to it.",
                "  - Set `code_system_used` to the value shown.",
                "  - Only use DISEASE_NAME for diagnoses not listed here.",
            ]
            for label, icd10, icd9 in coded_inclusions:
                if icd10 and icd9:
                    system = "MIXED"
                    detail = "ICD-10 present; ICD-9 also present (use ICD-10 only)"
                elif icd10:
                    system = "ICD-10"
                    detail = "ICD-10 codes provided"
                else:
                    system = "ICD-9"
                    detail = "ICD-9 codes only -- must be converted by requestor"
                code_lines.append(f"  label: \"{label}\" | {detail} | code_system_used='{system}'")
            provided_codes_note = "\n".join(code_lines)

    exclusion_note = ""
    if icd_table and icd_table.get("exclusions"):
        excl_lines = [
            "EXCLUSION CRITERIA DETECTED IN DATA REQUEST EXCLUSION TABLE (programmatically extracted):",
            "Populate request_summary.exclusion_criteria with each entry below.",
            "Each entry is a patient group the requestor explicitly wants EXCLUDED from the dataset.",
        ]
        for label, icd10, icd9 in icd_table["exclusions"]:
            codes    = icd10 or icd9
            code_str = f" [{', '.join(codes)}]" if codes else ""
            excl_lines.append(f"  - {label}{code_str}")
        exclusion_note = "\n".join(excl_lines)

    user_content = irb_content + request_content + clarification_content
    extra_notes  = "\n\n".join(n for n in [conflict_note, provided_codes_note, exclusion_note] if n)
    if extra_notes:
        user_content = user_content + [{"type": "text", "text": extra_notes}]

    if override_i2db:
        user_content = user_content + [{"type": "text", "text": (
            "REVIEWER OVERRIDE -- I2DB/RDC AUTHORIZATION CHECK WAIVED:\n"
            "The requirement to find a reference to I2DB, the Institute for Informatics, "
            "the Research Data Core (RDC), IRB 201607071, or Albert Lai is waived for this run "
            "at the explicit discretion of the data governance reviewer. Do NOT apply the "
            "GENOMIC DATA SOURCE AUTHORIZATION rules from the system prompt. Evaluate Tempus "
            "and genomic data elements purely on whether the IRB protocol itself authorizes them "
            "in scope and intent. If the protocol authorizes NGS or molecular profiling, treat "
            "that authorization as sufficient for this review."
        )}]
        print("  I2DB override: active")

    def _call():
        return client.messages.create(
            model=model,
            max_tokens=8192,
            system=SYSTEM_PROMPT,
            tools=[VALIDATION_TOOL],
            tool_choice={"type": "tool", "name": "submit_validation"},
            messages=[{"role": "user", "content": user_content}],
        )

    message     = _api_call_with_retry(_call)
    tool_blocks = [b for b in message.content if b.type == "tool_use"]
    if not tool_blocks:
        raise RuntimeError("API returned no tool_use block -- cannot extract result.")

    result = tool_blocks[0].input

    if precomputed_entries:
        cc = result.setdefault("consistency_checks", {})
        covered_exc_labels = {
            item[len("Exclusion -- "):]
            for entry in precomputed_entries
            for item in entry.get("conflicting_items", [])
            if item.startswith("Exclusion -- ")
        }
        kept = [
            e for e in cc.get("internal_request_conflicts", [])
            if not any(
                label in str(e.get("description", "")) or
                label in str(e.get("conflicting_items", ""))
                for label in covered_exc_labels
            )
        ]
        cc["internal_request_conflicts"] = precomputed_entries + kept

    if table_labels:
        for d in result.get("icd_code_analysis", {}).get("diagnoses_requested", []):
            if d.get("as_specified") in table_labels:
                d["source_table_label"] = d["as_specified"]

    if icd_table and icd_table.get("exclusions"):
        exact_exclusions = []
        for label, icd10, icd9 in icd_table["exclusions"]:
            codes = icd10 or icd9
            entry = label
            if codes:
                entry += f" [{', '.join(codes)}]"
            exact_exclusions.append(entry)
        result.setdefault("request_summary", {})["exclusion_criteria"] = exact_exclusions

    return result, message.usage, irb_format, request_format, clarification_format


def write_audit_log(
    irb_path: str,
    request_arg: str,
    irb_format: str,
    request_format: str,
    result: dict,
    usage,
    output_dir: Path,
    run_id: str,
    base_name: str = None,
    icd_resolution: dict = None,
    clarification_arg: str = None,
    clarification_format: str = None,
    override_i2db: bool = False,
) -> Path:
    inputs = {
        "irb_document":   str(irb_path),
        "irb_format":     irb_format,
        "request_source": (
            request_arg if len(request_arg) < 200 else request_arg[:200] + "..."
        ),
        "request_format": request_format,
    }
    if clarification_arg:
        inputs["clarification_source"] = (
            clarification_arg if len(clarification_arg) < 200 else clarification_arg[:200] + "..."
        )
        inputs["clarification_format"] = clarification_format
    if override_i2db:
        inputs["override_i2db"] = True

    audit = {
        "run_id":      run_id,
        "timestamp":   datetime.now(timezone.utc).isoformat() + "Z",
        "inputs":      inputs,
        "model":       os.environ.get("ANTHROPIC_MODEL", "claude-sonnet-4-6"),
        "token_usage": {
            "input_tokens":  usage.input_tokens,
            "output_tokens": usage.output_tokens,
        },
        "result": result,
    }
    if icd_resolution is not None:
        audit["icd_resolution"] = icd_resolution

    fname      = f"audit__{base_name}.json" if base_name else f"audit__{run_id}.json"
    audit_path = output_dir / fname
    with open(audit_path, "w") as f:
        json.dump(audit, f, indent=2)
    return audit_path


def print_report(result: dict):
    v      = result.get("validation", {})
    status = v.get("overall_status", "UNKNOWN")
    status_symbols = {"APPROVED": "[OK]", "DENIED": "[DENIED]", "NEEDS_REVIEW": "[REVIEW]"}
    symbol = status_symbols.get(status, "[?]")

    irb = result.get("irb_summary", {})
    req = result.get("request_summary", {})

    print(f"# IRB Validation Report -- {symbol} {status}\n")
    print(f"| | |")
    print(f"|---|---|")
    print(f"| **Protocol** | {irb.get('protocol_number', 'not found')} |")
    print(f"| **PI** | {irb.get('pi_name', 'not found')} |")
    print(f"| **Expires** | {irb.get('approval_expiry', 'not specified')} |")
    print(f"| **Requester** | {req.get('requester', 'not specified')} |")

    notes = result.get("reviewer_notes", "")
    if notes:
        print(f"\n## Summary for Reviewer\n")
        print(notes)
        print()

    icd        = result.get("icd_code_analysis", {})
    diagnoses  = icd.get("diagnoses_requested", [])
    icd_flags  = icd.get("icd_flags", [])
    if diagnoses:
        print(f"\n## ICD Code Analysis")
        code_systems    = [d.get("code_system_used", "") for d in diagnoses]
        uniform_system  = code_systems[0] if len(set(code_systems)) == 1 else None
        uniform_notes   = {
            "ICD-10":       "All diagnoses were provided by the requestor as ICD-10 codes.",
            "ICD-9":        "All diagnoses were provided as ICD-9 codes and must be converted before filtering.",
            "DISEASE_NAME": "All diagnoses were specified by disease name only -- no codes were provided in the data request. ICD-10 codes below are standard mappings that must be confirmed with the requestor.",
            "NONE":         "No codes or disease names were provided for any diagnosis.",
        }.get(uniform_system)
        if uniform_notes:
            print(f"*{uniform_notes}*\n")
        else:
            print(f"*Diagnoses listed in the data request and the ICD-10 codes that would be used to filter the dataset.*\n")

        for d in diagnoses:
            needs_flag = d.get("requires_clarification")
            marker     = "[!] " if needs_flag else ""
            print(f"### {marker}{d.get('as_specified', '?')}\n")
            if d.get("source_table_label"):
                print(f"*Source (DOCX table row): \"{d['source_table_label']}\"*")
            if not uniform_system:
                code_system  = d.get("code_system_used", "?")
                source_label = {
                    "ICD-10":       "Provided by requestor (ICD-10)",
                    "ICD-9":        "Provided by requestor (ICD-9 -- must be converted)",
                    "MIXED":        "Provided by requestor (mix of ICD-9 and ICD-10)",
                    "DISEASE_NAME": "Specified by disease name only -- no codes provided",
                    "NONE":         "No codes or disease name provided",
                }.get(code_system, code_system)
                print(f"- **How it was specified (in the data request):** {source_label}")
            if d.get("icd9_codes"):
                print(f"- **ICD-9 codes (from data request, cannot be used directly):** {', '.join(d['icd9_codes'])}")
            if d.get("icd10_codes"):
                codes = ", ".join(d["icd10_codes"])
                auto  = "[auto-resolved] " if "Auto-resolved" in d.get("mapping_notes", "") else ""
                print(f"- **ICD-10 codes to use for filtering** {auto}-- `{codes}`")
            if d.get("qualifier_codes"):
                quals = ", ".join(d["qualifier_codes"])
                print(f"- **Additional qualifier codes** (e.g. receptor status, metastatic): `{quals}`")
            if d.get("combination_logic"):
                print(f"- **Recommended filter logic:**\n  ```\n  {d['combination_logic']}\n  ```")
            if d.get("mapping_notes"):
                notes_text = d["mapping_notes"].replace("[Auto-resolved from ICD-10 reference -- verify before use] ", "")
                print(f"- **Notes:** {notes_text}")
            if needs_flag:
                print(f"- [!] **Action required before this diagnosis can be filtered.**")
            print()

    excl_diagnoses = icd.get("exclusion_diagnoses", [])
    if excl_diagnoses:
        print(f"### Exclusion Diagnoses\n")
        print(f"*ICD codes used to exclude patients from the cohort, as specified in the data request.*\n")
        for d in excl_diagnoses:
            print(f"**{d.get('as_specified', '?')}**")
            if d.get("icd9_codes"):
                print(f"- **ICD-9 (cannot be used directly):** {', '.join(d['icd9_codes'])}")
            if d.get("icd10_codes"):
                print(f"- **ICD-10 exclusion codes:** `{', '.join(d['icd10_codes'])}`")
            if d.get("qualifier_condition"):
                print(f"- **Condition:** {d['qualifier_condition']}")
            if d.get("notes"):
                print(f"- **Notes:** {d['notes']}")
            print()

    if icd_flags:
        print(f"**ICD issues requiring attention:**\n")
        for flag in icd_flags:
            print(f"- [!] {flag}")
        print()

    personnel = v.get("personnel_check", {})
    pstatus   = personnel.get("status", "")
    psym      = {"APPROVED": "[OK]", "DENIED": "[DENIED]", "NOT_SPECIFIED": "[!]", "NEEDS_REVIEW": "[!]"}.get(pstatus, "[?]")
    print(f"## Personnel -- {psym} {pstatus}")
    print(f"*Checks whether everyone who signed the data request is listed as approved personnel on the IRB.*\n")
    print(f"{personnel.get('notes', '')}\n")

    approved = v.get("approved_elements", [])
    if approved:
        print(f"## [OK] Data Elements -- Approved ({len(approved)})")
        print(f"*These items were requested and are explicitly covered by the IRB protocol.*\n")
        for el in approved:
            print(f"**{el['element']}**  ")
            print(f"{el['rationale']}\n")

    denied = v.get("denied_elements", [])
    if denied:
        print(f"## [DENIED] Data Elements -- Denied ({len(denied)})")
        print(f"*These items were requested but fall outside what the IRB has approved. They cannot be released.*\n")
        for el in denied:
            print(f"**{el['element']}**  ")
            print(f"{el['rationale']}\n")

    ambiguous = v.get("ambiguous_elements", [])
    if ambiguous:
        print(f"## [!] Data Elements -- Needs Human Review ({len(ambiguous)})")
        print(f"*These items could not be clearly approved or denied and require a reviewer to make a judgment call.*\n")
        for el in ambiguous:
            print(f"**{el['element']}**  ")
            print(f"*Issue:* {el['reason']}  ")
            if el.get("irb_language") and el["irb_language"] != "N/A":
                print(f"*Relevant IRB text:* \"{el['irb_language']}\"  ")
            if el.get("recommendation"):
                print(f"*What to do:* {el['recommendation']}")
            print()

    pop_checks = v.get("population_filter_check", [])
    if pop_checks:
        print(f"## Population Scope Check")
        print(f"*Does the group of patients described in the data request match what the IRB approved?*\n")
        print("| Patient filter | Status | Notes |")
        print("|----------------|--------|-------|")
        for f in pop_checks:
            sym          = {"WITHIN_SCOPE": "[OK]", "OUTSIDE_SCOPE": "[OUT]", "AMBIGUOUS": "[?]"}.get(f["status"], "?")
            status_label = {"WITHIN_SCOPE": "[OK] OK", "OUTSIDE_SCOPE": "[OUT] Out of scope", "AMBIGUOUS": "[?] Unclear"}.get(f["status"], f["status"])
            notes_cell   = f["notes"].replace("|", "\\|").replace("\n", " ")
            print(f"| {f['filter']} | {status_label} | {notes_cell} |")
        print()

    available = result.get("available_but_not_requested", [])
    if available:
        print(f"## IRB Scope Not Used")
        print(f"*These elements or populations are explicitly approved by the IRB but were not included in the data request.*\n")
        for item in available:
            print(f"- {item}")
        print()

    cc       = result.get("consistency_checks", {})
    internal = cc.get("internal_request_conflicts", [])
    cross    = cc.get("cross_document_conflicts", [])
    if internal or cross:
        print(f"## [!] Conflicts Found")
        if internal:
            print(f"### Problems within the data request itself\n")
            print(f"*The following issues were found inside the data request document.*\n")
            for c in internal:
                print(f"**{c['description']}**  ")
                if c.get("conflicting_items"):
                    print(f"*Conflicting items:* {', '.join(c['conflicting_items'])}  ")
                if c.get("recommendation"):
                    print(f"*What to do:* {c['recommendation']}")
                print()
        if cross:
            print(f"### Conflicts between the IRB protocol and the data request\n")
            print(f"*The data request asks for something that the IRB does not permit, or vice versa.*\n")
            for c in cross:
                print(f"**{c['description']}**  ")
                if c.get("irb_basis") and c["irb_basis"] != "N/A":
                    print(f"*What the IRB says:* \"{c['irb_basis']}\"  ")
                if c.get("request_basis"):
                    print(f"*What the request says:* \"{c['request_basis']}\"  ")
                if c.get("recommendation"):
                    print(f"*What to do:* {c['recommendation']}")
                print()


def _extract_flag(argv: list, flag: str) -> tuple:
    value   = None
    cleaned = []
    i       = 0
    while i < len(argv):
        if argv[i] == flag and i + 1 < len(argv):
            value = argv[i + 1]
            i    += 2
        else:
            cleaned.append(argv[i])
            i += 1
    return cleaned, value


def main():
    usage_msg = f"""Usage:
  venv/bin/python data-claim.py <irb_document> <request_document> [options]

Arguments:
  irb_document        IRB protocol file
  request_document    Data request file or inline text

Options:
  --clarification <path_or_text>
                      Follow-up clarifications resolving ambiguous wording in the data
                      request. Can be a file path or inline text. DOCX clarifications
                      with ICD tables are merged with the request's tables -- clarification
                      entries supersede request entries for the same diagnosis label.
                      All scope must still fall within IRB limits.
  --icd-file <path>   ICD-10 reference CSV. Defaults to the newest
                      ICD10_codes/icd10cm-oncology-*.csv found next to this script.
                      When a disease name appears without ICD-10 codes, a second Claude
                      API call searches this CSV and returns the matching codes.
  --override-i2db     Waive the I2DB/RDC authorization check at reviewer discretion.
                      Tempus genomic data will be evaluated on IRB scope alone rather
                      than auto-denied for missing I2DB/RDC protocol language.
  --output-dir <path> Directory to write report and audit JSON into. Defaults to the
                      same directory as the IRB document. Created if it does not exist.

Supported formats: {SUPPORTED_FORMATS}

Output is written to --output-dir (or next to the IRB document if not specified):
  report__*.md    -- human-readable compliance report
  audit__*.json   -- full structured result (input to excavator/orchestrator.py)
"""

    argv          = sys.argv[1:]
    override_i2db = "--override-i2db" in argv
    argv          = [a for a in argv if a != "--override-i2db"]
    argv, icd_file          = _extract_flag(argv, "--icd-file")
    argv, clarification_arg = _extract_flag(argv, "--clarification")
    argv, output_dir_arg    = _extract_flag(argv, "--output-dir")

    if icd_file is None:
        candidates = sorted(Path(__file__).parent.glob("ICD10_codes/icd10cm-oncology-*.csv"))
        if candidates:
            icd_file = str(candidates[-1])

    if len(argv) != 2:
        print(usage_msg)
        sys.exit(1)

    irb_path    = Path(argv[0])
    request_arg = argv[1]
    study_dir   = irb_path.parent
    output_dir  = Path(output_dir_arg) if output_dir_arg else study_dir
    output_dir.mkdir(parents=True, exist_ok=True)
    run_id      = datetime.now(timezone.utc).strftime("%Y%m%d_%H%M%S")

    print(f"\nValidating data request against IRB...")
    print(f"  IRB path: {irb_path}")

    result, usage, irb_format, request_format, clarification_format = validate_request(
        str(irb_path), request_arg, clarification_arg=clarification_arg,
        override_i2db=override_i2db,
    )

    icd_resolution = None
    if icd_file:
        names_to_resolve = [
            d["as_specified"]
            for d in result.get("icd_code_analysis", {}).get("diagnoses_requested", [])
            if d.get("requires_clarification") and d.get("code_system_used") == "DISEASE_NAME"
        ]
        if names_to_resolve:
            print(f"  Resolving ICD-10 codes for: {names_to_resolve}")
            try:
                lookup_result, lookup_usage = resolve_icd_codes(
                    names_to_resolve, icd_file, irb_context=result.get("irb_summary", {})
                )
                result         = apply_icd_resolution(result, lookup_result)
                icd_resolution = {
                    "icd_file":    str(icd_file),
                    "lookup_result": lookup_result,
                    "token_usage": {
                        "input_tokens":  lookup_usage.input_tokens,
                        "output_tokens": lookup_usage.output_tokens,
                    },
                }
                print(
                    f"  ICD resolution tokens: "
                    f"{lookup_usage.input_tokens} in / {lookup_usage.output_tokens} out"
                )
            except Exception as e:
                print(f"  WARNING: ICD resolution failed: {e}")

    irb_stem   = irb_path.stem
    req_stem   = Path(request_arg).stem if Path(request_arg).exists() else "inline"
    base_name  = f"{study_dir.name}__{irb_stem}__{req_stem}__{run_id}"

    audit_path = write_audit_log(
        str(irb_path), request_arg, irb_format, request_format,
        result, usage, output_dir, run_id, base_name,
        icd_resolution=icd_resolution,
        clarification_arg=clarification_arg,
        clarification_format=clarification_format,
        override_i2db=override_i2db,
    )

    import io, contextlib
    buf = io.StringIO()
    with contextlib.redirect_stdout(buf):
        print_report(result)
    report_path = output_dir / f"report__{base_name}.md"
    report_path.write_text(buf.getvalue())

    print_report(result)
    print(f"Report:    {report_path}")
    print(f"Audit log: {audit_path}")
    print(f"Tokens:    {usage.input_tokens} in / {usage.output_tokens} out")


if __name__ == "__main__":
    main()
